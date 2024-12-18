from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Ethical Hacking Questions and Answers', 0)

# Adding the content
questions_and_answers = [
    ("What are Hackers & Types of Hacking & Its Phases?", """
    **Answer:**
    **Hackers** are individuals who possess the technical skills to gain unauthorized access to systems, networks, or data, often with malicious intent or for personal gain. However, not all hackers have bad intentions, as some work ethically to protect systems.

    - **Types of Hackers:**
      1. **White Hat Hackers (Ethical Hackers):** These hackers use their skills for ethical purposes, such as finding vulnerabilities and helping organizations protect their systems. They work with the permission of the organization.
      2. **Black Hat Hackers:** These are malicious hackers who exploit vulnerabilities for personal or financial gain, causing harm to individuals or organizations. They often engage in illegal activities such as stealing data, deploying malware, or disrupting services.
      3. **Gray Hat Hackers:** These hackers fall between the ethical and unethical categories. They may find vulnerabilities without permission and report them, but their methods may not always be ethical.

    **Phases of Hacking:**
    1. **Reconnaissance (Information Gathering):** Hackers gather as much information as possible about the target, such as network configurations, domain names, IP addresses, and employee details.
    2. **Scanning:** In this phase, hackers use tools to discover potential vulnerabilities in the target system.
    3. **Gaining Access:** This involves exploiting a vulnerability to gain unauthorized access to the system. Methods include SQL injection, buffer overflows, or exploiting weak passwords.
    4. **Maintaining Access:** Once access is gained, hackers ensure they can re-enter the system, often by installing backdoors or creating new user accounts.
    5. **Clearing Tracks:** The final phase involves erasing logs, deleting traces of the attack, or manipulating system files to cover up their activities.
    """),
    ("What is SQL Injection? Provide an example.", """
    **Answer:**
    **SQL Injection** is a type of attack where a malicious SQL query is inserted into an input field, which is then executed by the database, allowing the attacker to manipulate the database in unintended ways. This attack typically occurs when user inputs are not properly validated or sanitized.

    **Example of SQL Injection:**
    Consider a login form where the username and password fields are used to authenticate a user:
    - **Login Query:**  
      ```sql
      SELECT * FROM users WHERE username = 'user_input' AND password = 'password_input';
      ```

    If the input fields are not sanitized, an attacker could input the following into the username or password fields:
    - **Username:** `' OR 1=1 --`
    - **Password:** `anything`

    This would modify the SQL query to:
    ```sql
    SELECT * FROM users WHERE username = '' OR 1=1 -- AND password = 'anything';
    ```
    Here, the `OR 1=1` condition is always true, allowing the attacker to bypass the authentication and gain unauthorized access.

    **Consequences of SQL Injection**:
    - Unauthorized access to sensitive data.
    - Deletion or modification of data.
    - Potential administrative access to the database.
    """),
    ("What is the difference between a Virus, Worm, and Trojan?", """
    **Answer:**
    A **Virus**, **Worm**, and **Trojan** are all types of malware, but they have distinct characteristics:

    1. **Virus:**
       - A **virus** is a piece of malicious code that attaches itself to a legitimate program or file. Once executed, it can spread by infecting other files or programs on the system. It requires a host program to run.
       - **Example**: A virus could be embedded in an email attachment that, once opened, infects the system.

    2. **Worm:**
       - A **worm** is a self-replicating program that spreads across networks without requiring a host file. Unlike a virus, a worm can spread independently and is designed to exploit network vulnerabilities.
       - **Example**: A worm may infect a system through a vulnerability in an email system or network service and then replicate itself across other systems on the network.

    3. **Trojan (Trojan Horse):**
       - A **Trojan** is a type of malware that masquerades as legitimate software to trick the user into installing it. Unlike viruses and worms, Trojans do not replicate themselves. They can give attackers backdoor access to the system.
       - **Example**: A Trojan may appear as a fake software update that, when executed, installs a backdoor for hackers to access the system remotely.

    **Key Differences:**
    - A **virus** attaches itself to programs, a **worm** spreads through networks, and a **Trojan** deceives users by appearing legitimate.
    """),
    ("What is a Dos Attack & Its Types? What is DDoS?", """
    **Answer:**
    A **Denial of Service (DoS)** attack is an attempt to make a system or network service unavailable to users by overwhelming it with a flood of traffic or requests. DoS attacks can disrupt services and cause financial or reputational damage.

    **Types of DoS Attacks**:
    1. **Flooding Attacks**: These attacks involve sending massive amounts of traffic to the target server, making it unable to respond to legitimate requests. Common examples include:
       - **Ping of Death**: Sending oversized ping packets to overwhelm the target.
       - **SYN Flood**: Exploiting the TCP handshake process by sending SYN requests without completing the handshake, causing the target to exhaust resources.
    
    2. **Resource Exhaustion Attacks**: These attacks consume system resources, such as CPU, memory, or disk space, rendering the system unusable.
       - **Example**: Sending a large number of requests that consume all available resources on the server, slowing down or crashing the system.

    **Distributed Denial of Service (DDoS)**:
    A **DDoS** attack is a type of DoS attack where the attack is launched from multiple compromised systems, often called a botnet. The distributed nature of the attack makes it more difficult to defend against.
    - **Example**: A large botnet of infected computers floods a target website with excessive traffic, causing it to crash.
    """),
    ("What is Steganography and its types and applications?", """
    **Answer:**
    **Steganography** is the practice of concealing information within other non-suspicious data (such as images, audio files, or video) to prevent detection. Unlike encryption, which transforms data into an unreadable form, steganography hides the existence of the message itself.

    **Types of Steganography**:
    1. **Image Steganography**: Information is hidden in the least significant bits (LSBs) of an image's pixels. Changes to these bits are typically not visible to the human eye.
       - **Example**: A hidden message can be embedded into an image file, making it appear unchanged while secretly containing data.
    
    2. **Audio Steganography**: Data is embedded in audio files, either by altering the amplitude or phase of the audio waves or by hiding data within frequencies that are inaudible to humans.
       - **Example**: Hiding a secret message within an audio recording of a song.

    3. **Video Steganography**: Hiding data within video frames or audio tracks that accompany the video. This technique can be more complex due to the larger size of video files.
       - **Example**: Embedding a message in the silent parts of a video.

    **Applications**:
    - **Digital Watermarking**: Protecting copyrights by embedding a unique identifier in media files.
    - **Covert Communication**: Sending secret messages without detection, often used by spies or in confidential communications.
    - **Data Integrity**: Ensuring the integrity of transmitted data by hiding checksum information in the file.
    """),
    ("What is Session Hijacking?", """
    **Answer:**
    **Session Hijacking** is an attack where an attacker takes control of a user's active session, typically by stealing or predicting the session token or session ID used for authentication. Once the attacker hijacks the session, they can impersonate the legitimate user and gain unauthorized access to sensitive information or services.

    **How Session Hijacking Works:**
    - **Session Fixation**: The attacker sets the victim’s session ID to one they control before the victim logs in.
    - **Session Sidejacking**: The attacker intercepts network traffic containing the session cookie, typically over an unsecured Wi-Fi network.
    - **Session Sniffing**: The attacker listens to network traffic to obtain the session ID, usually when session IDs are sent in unencrypted HTTP requests.

    **Consequences**:
    - Unauthorized access to personal accounts, financial data, or sensitive systems.
    - The attacker can perform actions as the victim, including making transactions or accessing private files.
    """)
]

# Loop to add content to the document
for question, answer in questions_and_answers:
    doc.add_heading(question, level=1)
    doc.add_paragraph(answer)

# Save the document
file_path = "C:/Users/avina/OneDrive/Desktop/ethical_hacking_questions_and_answers.docx"
doc.save(file_path)

file_path
